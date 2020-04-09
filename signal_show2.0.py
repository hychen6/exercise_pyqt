from PyQt5.QtWidgets import QApplication,QMainWindow,QHBoxLayout,QFileDialog,qApp,QToolTip,QStatusBar,QMenu,QWidget,QTabWidget,QTabBar,QVBoxLayout
from PyQt5 import QtWidgets
from PyQt5.QtCore import Qt
from PyQt5 import QtCore
import os
from PyQt5.QtWidgets import QTableWidgetItem
import sys
import xlwt
import signal_show_main
import pyqtgraph as pg
import pyqtgraph.exporters
import numpy as np
from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches
curPath = os.path.abspath(os.path.dirname(__file__))
rootPath = os.path.split(curPath)[0]
sys.path.append(rootPath)
import matplotlib as mpl
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D

from matplotlib.pyplot import MultipleLocator
from matplotlib.backends.backend_qt5agg import (FigureCanvas,
                NavigationToolbar2QT as NavigationToolbar)







class initUI(QMainWindow):
    def __init__(self,parent=None):
        super().__init__()
        self.initUi()

    def initUi(self):

        global fileName1
        fileName1=None
        self.mainwindow=signal_show_main.Ui_MainWindow()
        self.mainwindow.setupUi(self)
#        self.mainwindow.pushButton.clicked.connect(self.new)
        self.mainwindow.actionopen.triggered.connect(self.newopen)
        self.mainwindow.actionquit.triggered.connect(qApp.quit)
        self.mainwindow.pushButton.clicked.connect(self.open_dir)
        self.mainwindow.actionasave.triggered.connect(self.SavePictureTo)
#        self.mainwindow.verticalLayout_2.setContextMenuPolicy(Qt.CustomContextMenu)
       # self.mainwindow.tab.customContextMenuRequested.connect(self.custom_right_menu1)
#        self.mainwindow.verticalLayout_3.setContextMenuPolicy(Qt.CustomContextMenu)
        #self.mainwindow.tab_3.customContextMenuRequested.connect(self.custom_right_menu)
        self.mainwindow.actionreport.triggered.connect(self.report)
        self.mainwindow.actionexcel.triggered.connect(self.excel)
        self.mainwindow.actionzdata.setChecked(True)
        self.mainwindow.actionpitch.setChecked(True)
        self.mainwindow.actionindextable.setChecked(True)
        self.mainwindow.action3d.setChecked(True)
        #str="QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setStyleSheet(str)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2))
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_3))
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_5))
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_6))
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2), False)
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_3),False)
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_5), False)
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_6), False)


        self.mainwindow.actionzdata.toggled.connect(self.removezdata)
        self.mainwindow.actionhdata.toggled.connect(self.addhdata)
        self.mainwindow.actionbdata.toggled.connect(self.addbdata)
        self.mainwindow.actionpitch.toggled.connect(self.removepitch)
        self.mainwindow.actionroll.toggled.connect(self.addroll)
        self.mainwindow.actionheading.toggled.connect(self.addheading)
        self.mainwindow.actionindextable.toggled.connect(self.removeindextable)
        self.mainwindow.action3d.toggled.connect(self.remove3dplot)
        self.mainwindow.actionsave.triggered.connect(self.savepicture)
        self.widget = None
        self.widget1 = None
        self.widget2 = None
        self.widget3 = None
        self.widget4 = None
        self.widget5 = None
        self.widget8=None
        self.exedir=os.getcwd()
        #self.tab_5=None

        #if self.mainwindow.actionhdata.isChecked()==True:
            #self.mainwindow.actionhdata.toggled.connect(self.removehdata)
        #else:
    def open_dir(self):
        self.dir_path = QFileDialog.getExistingDirectory(self, "choose directory")
        if not os.path.exists(self.dir_path):
            return
        #self.dir_path = self.dir_path.replace('/', '\\')  # windows下需要进行文件分隔符转换
        self.mainwindow.lineEdit.setText(self.dir_path)
        self.path_list=os.listdir(self.dir_path)
        self.exedir=os.getcwd()
          # 获取该路径下所有的文件以及目录并显示在listwidget中
        #self.path_list = mylib.getAllPath(self.dir_path)
          # print("path_list",self.path_list)
        self.mainwindow.listWidget.clear()
        self.mainwindow.listWidget.itemClicked.connect(self.listnew)
        #if len(self.path_list) > 0:
         #   self.mainwindow.listWidget.addItems(self.path_list)
         #   fileName1=
        for i in range(0,len(self.path_list)):
            self.mainwindow.listWidget.addItem(self.path_list[i])
            #fileName1 =self.path_list[i]



            #self.toolbar_run.setCheckable(True)



        #self.mainwindow.verticalLayout.addWidget(self.widget)
    def listnew(self):
        global fileName1
        currentrow=self.mainwindow.listWidget.currentItem()
        if (os.path.splitext(currentrow.text())[1] in['.fmd'] )==True:
            fileName1=currentrow.text()
        if self.dir_path:
            os.chdir(self.dir_path)
        self.new()
        os.chdir(self.exedir)
    def addzdata(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_1), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_1,'磁垂直分量')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_1)
        self.mainwindow.actionzdata.toggled.connect(self.removezdata)
    def addhdata(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_2,'磁水平分量')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_2)
        self.mainwindow.actionhdata.toggled.connect(self.removehdata)
    def addbdata(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_3), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_3,'磁总量')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_3)
        self.mainwindow.actionbdata.toggled.connect(self.removebdata)
    def addpitch(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_4), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_4,'罗盘俯仰角')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_4)
        self.mainwindow.actionpitch.toggled.connect(self.removepitch)
    def addroll(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_5), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_5,'罗盘横滚角')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_5)
        self.mainwindow.actionroll.toggled.connect(self.removeroll)
    def addheading(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_6), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_6,'罗盘方位角')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_6)
        self.mainwindow.actionheading.toggled.connect(self.removeheading)
    def addindextable(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_7), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab_7,'数据表格')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab_7)
        self.mainwindow.actionindextable.toggled.connect(self.removeindextable)
    def add3dplot(self):
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2), True)
        self.mainwindow.tabWidget.addTab(self.mainwindow.tab,'运动轨迹')
        self.mainwindow.tabWidget.setCurrentWidget(self.mainwindow.tab)
        self.mainwindow.action3d.toggled.connect(self.remove3dplot)


    def removebdata(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_3))
        self.mainwindow.actionbdata.toggled.connect(self.addbdata)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removehdata(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tab_3.close()
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_2))
        self.mainwindow.actionhdata.toggled.connect(self.addhdata)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removezdata(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_1), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_1))
        self.mainwindow.actionzdata.toggled.connect(self.addzdata)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removepitch(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_4), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_4))
        self.mainwindow.actionpitch.toggled.connect(self.addpitch)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removeroll(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_5), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_5))
        self.mainwindow.actionroll.toggled.connect(self.addroll)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removeheading(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_6), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_6))
        self.mainwindow.actionheading.toggled.connect(self.addheading)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def removeindextable(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_7), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_7))
        self.mainwindow.actionindextable.toggled.connect(self.addindextable)
        #self.mainwindow.tabWidget.setStyleSheet(str)
    def remove3dplot(self):
        #str = "QTabBar::tab:disabled {width: 0; color: transparent;}"
        #self.mainwindow.tabWidget.setTabEnabled(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab_7), False)
        self.mainwindow.tabWidget.removeTab(self.mainwindow.tabWidget.indexOf(self.mainwindow.tab))
        self.mainwindow.action3d.toggled.connect(self.add3dplot)
        #self.mainwindow.tabWidget.setStyleSheet(str)


    def gettable(self):
        self.mainwindow.tableWidget.setColumnCount(7)
        self.mainwindow.tableWidget.setRowCount(len(self.deep))
        columnname = ['深度', '磁垂直分量', '磁水平分量', '磁总量', '罗盘俯仰角','横滚角','方位角']
        self.mainwindow.tableWidget.setHorizontalHeaderLabels(columnname)
        for i in range(len(self.deep)):
            self.mainwindow.tableWidget.setItem(i, 0, QTableWidgetItem(str(self.deep[i])))
            self.mainwindow.tableWidget.setItem(i, 1, QTableWidgetItem(str(self.zdata[i])))
            self.mainwindow.tableWidget.setItem(i, 2, QTableWidgetItem(str(self.hdata[i])))
            self.mainwindow.tableWidget.setItem(i, 3, QTableWidgetItem(str(self.bdata[i])))
            self.mainwindow.tableWidget.setItem(i, 4, QTableWidgetItem(str(self.pitch[i])))
            self.mainwindow.tableWidget.setItem(i, 5, QTableWidgetItem(str(self.roll[i])))
            self.mainwindow.tableWidget.setItem(i, 6, QTableWidgetItem(str(self.headind[i])))
    def report(self):
        if not fileName1:
            return
        template = "RSM-RLT(A)钢筋笼长度检测报告.docx"
        document_1 = MailMerge(template)
        document_1.merge(
            project_name=self.ProjectName,
            construct_institution=self.ConstrucInstitution,
            test_institution=self.TestInstitution,
            tester_name=self.TesterName,
            test_date=self.TestDate,
        )
        #document_1.add_picture
#        document_1.add_picture('test_zdata.png')
        document_1.write('file.docx')
        doc=Document('file.docx')
        self.savepicture()
        run = doc.tables[2].cell(0, 0).paragraphs[0].add_run()
        run.add_picture('test_zdata.png',width=Inches(5.0))
        run = doc.tables[2].cell(1, 0).paragraphs[0].add_run()
        run.add_picture('test_hdata.png',width=Inches(5.0))
        run = doc.tables[2].cell(2, 0).paragraphs[0].add_run()
        run.add_picture('test_bdata.png',width=Inches(5.0))
        run = doc.tables[2].cell(3, 0).paragraphs[0].add_run()
        run.add_picture('test_pitch.png',width=Inches(5.0))
        run = doc.tables[2].cell(4, 0).paragraphs[0].add_run()
        run.add_picture('test_roll.png',width=Inches(5.0))
        run = doc.tables[2].cell(5, 0).paragraphs[0].add_run()
        run.add_picture('test_heading.png',width=Inches(5.0))
        fileName2, ok2 = QFileDialog.getSaveFileName(self,
                                                     "文件保存",
                                                     "C:/",
                                                     "Word Files (*.docx)")

        if fileName2:
            doc.save(fileName2)
    def excel(self):
        if not fileName1:
            return
        workbook = xlwt.Workbook(encoding='utf-8')  # 新建工作簿
        sheet1 = workbook.add_sheet("检测数据")  # 新建sheet
        sheet1.write(0, 0, "深度")  # 第1行第1列数据
        sheet1.write(0, 1, "垂直磁分量")  # 第1行第2列数据
        sheet1.write(0, 2, "水平磁分量")
        sheet1.write(0, 3, "磁总量")
        sheet1.write(0, 4, "罗盘俯仰角")
        sheet1.write(0, 5, "横滚角")
        sheet1.write(0, 6, "方位角")
        for i in range(1, len(self.deep) + 1):
            sheet1.write(i, 0, self.deep[i-1])  # 第1行第1列数据
            sheet1.write(i, 1, self.zdata[i-1])  # 第1行第2列数据
            sheet1.write(i, 2, self.hdata[i-1])
            sheet1.write(i, 3, self.bdata[i-1])
            sheet1.write(i, 4, self.pitch[i-1])
            sheet1.write(i, 5, self.roll[i-1])
            sheet1.write(i, 6, self.headind[i-1])
        fileName3, ok3 = QFileDialog.getSaveFileName(self,
                                                     "文件保存",
                                                     "C:/",
                                                     "Excel Files (*.xls)")
        if fileName3:

            workbook.save(fileName3)
    def savepicture(self):
        if not fileName1:
            return
        self.ex_zdata.export(fileName="test_zdata.png")
        self.ex_hdata.export(fileName="test_hdata.png")
        self.ex_bdata.export(fileName="test_bdata.png")
        self.ex_pitch.export(fileName="test_pitch.png")
        self.ex_roll.export(fileName="test_roll.png")
        self.ex_heading.export(fileName="test_heading.png")
    def SavePictureTo(self):
        if not fileName1:
            return
        fileName3= QFileDialog.getExistingDirectory(self, "choose directory")
        if not os.path.exists(fileName3):
            return
        if fileName3:
            os.chdir(fileName3)
        self.ex_zdata.export(fileName="test_zdata.png")
        self.ex_hdata.export(fileName="test_hdata.png")
        self.ex_bdata.export(fileName="test_bdata.png")
        self.ex_pitch.export(fileName="test_pitch.png")
        self.ex_roll.export(fileName="test_roll.png")
        self.ex_heading.export(fileName="test_heading.png")
        os.chdir(self.exedir)
    def newopen(self):
        self.getfile()
        self.new()
    def new(self):
        if fileName1:
            self.plot = PlotWidget(self)

            self.VerHard=str(self.plot.VerHard)   # 硬件版本号，例如19A3表示2019年10(对应A)月份第3版
            self.VerSoft = str(self.plot.VerSoft) # 软件版本号，同上
            self.ProjectName =str(self.plot.ProjectName)  # 工程名称
            self.ConstrucInstitution = str(self.plot.ConstrucInstitution) # 施工单位
            self.TestInstitution = str(self.plot.TestInstitution) # 检测执行单位名称
            self.TesterName = str(self.plot.TestInstitution) # 检测员姓名(编号)
            self.TestDate = str(self.plot.TestDate)  # 检测时间
            self.DesignLength = str(self.plot.DesignLength) # 钢筋笼设计长度，单位m（四舍五入取整）
            self.TestPointNo = str(self.plot.TestPointNo)  # 测点编号，对应孔桩编号
            self.InitialDepth = str(self.plot.InitialDepth)  # 初始深度，单位cm
            self.PileDiameter = str(self.plot.PileDiameter)  # 基桩的直径，单位cm
            self.Orientation = str(self.plot.Orientation) # 测孔相对桩的方位偏角，单位度
            self.Range = str(self.plot.Range) # 桩和测孔的最小距离，单位cm
            self.StepDistance = str(self.plot.StepDistance) # 检测步距： 5、10、15、20、15, 单位cm
            self.TestMode = str(self.plot.TestMode)# 采集方式，手动0；自动1
            self.MoveDirection = str(self.plot.MoveDirection)
            ## 深度值（4个字节） + 磁垂直Z分量+磁水平H分量+磁B总量 +罗盘Pitch俯仰角+ Roll横滚角+ Heading方位角 +0x0D+ 0x0A
            self.deep=self.plot.a
            self.zdata = self.plot.zdata
            self.hdata = self.plot.hdata
            self.bdata = self.plot.bdata
            self.pitch = self.plot.pitch
            self.roll = self.plot.roll
            self.headind = self.plot.headind
            self.gettable()
            self.getzdataplot()
            self.gethdataplot()
            self.getbdataplot()
            self.getpitchplot()
            self.getrollplot()
            self.getheadingplot()
            self.getplot3d()
            #self.plot3d=Plot3d(self.mainwindow.tab)
            #self.mainwindow.tabWidget.addTab(self.mainwindow.tab, '运动轨迹')
            self.mainwindow.actionsave.setEnabled(True)
            self.mainwindow.actionasave.setEnabled(True)
            self.mainwindow.actionreport.setEnabled(True)
            self.mainwindow.actionexcel.setEnabled(True)
    def getplot3d(self):
        self.plot3d=Plot3d(self)
        #self.plot3d.test()
        self.mainwindow.horizontalLayout_3.removeWidget(self.widget8)
        self.mainwindow.horizontalLayout_3.addWidget(self.plot3d.figCanvas)
        self.widget8 = self.plot3d.figCanvas

    def getzdataplot(self):
        self.plot = PlotWidget(self)
        self.plot.s = self.zdata
        self.plot.label_left1 = "磁场垂直分量曲线 Z(mGs)"
        self.plot.label_left2 = "磁场垂直分量梯度曲线 dZ/dh(mGs/m)"
        self.plot.plot()
        self.ex_zdata = self.plot.ex
        self.mainwindow.verticalLayout1.removeWidget(self.widget)
        self.mainwindow.verticalLayout1.addWidget(self.plot.win)
        self.widget = self.plot.win
    def gethdataplot(self):

        self.plot3 = PlotWidget(self)
        self.plot3.s = self.hdata
        self.plot3.label_left1="磁场水平分量曲线 X(mGs)"
        self.plot3.label_left2= "磁场水平分量梯度曲线 dX/dh(mGs/h)"
        self.plot3.plot()
        self.mainwindow.verticalLayout2.removeWidget(self.widget1)
        self.mainwindow.verticalLayout2.addWidget(self.plot3.win)
        self.widget1 = self.plot3.win
        self.ex_hdata=self.plot3.ex
    def getbdataplot(self):

        self.plot4 = PlotWidget(self)
        self.plot4.s = self.bdata
        self.plot4.label_left1 = "磁场总量曲线 B(mGs)"
        self.plot4.label_left2 = "磁场总量梯度曲线 dB/dh(mGs/h)"
        self.plot4.plot()
        self.mainwindow.verticalLayout3.removeWidget(self.widget2)
        self.mainwindow.verticalLayout3.addWidget(self.plot4.win)
        self.widget2 = self.plot4.win
        self.ex_bdata = self.plot4.ex
    def getpitchplot(self):
        self.plot2 = PlotWidget(self)
        self.plot2.s = self.pitch
        self.plot2.label_left1 = "罗盘俯仰角曲线"
        self.plot2.label_left2 = "罗盘俯仰角梯度曲线"
        self.plot2.plot()
        self.ex_pitch = self.plot2.ex
        self.mainwindow.verticalLayout4.removeWidget(self.widget5)
        self.mainwindow.verticalLayout4.addWidget(self.plot2.win)
        self.widget5 = self.plot2.win
    def getrollplot(self):


        self.plot5 = PlotWidget(self)
        self.plot5.s = self.roll
        self.plot5.label_left1 = "罗盘横滚角曲线"
        self.plot5.label_left2="罗盘横滚角梯度曲线"
        self.plot5.plot()
        self.mainwindow.verticalLayout5.removeWidget(self.widget3)
        self.mainwindow.verticalLayout5.addWidget(self.plot5.win)
        self.widget3 = self.plot5.win
        self.ex_roll=self.plot5.ex


    def getheadingplot(self):


        self.plot6 = PlotWidget(self)
        self.plot6.s = self.hdata
        self.plot6.label_left1 = "罗盘方位角曲线"
        self.plot6.label_left2="罗盘方位角梯度曲线"
        self.plot6.plot()
        self.mainwindow.verticalLayout6.removeWidget(self.widget4)
        self.mainwindow.verticalLayout6.addWidget(self.plot6.win)
        self.widget4 = self.plot6.win
        self.ex_heading=self.plot6.ex


#            self.mainwindow.tabWidget.removeTab(1)

    def getfile(self):
        global fileName1
        fileName, filetype = QFileDialog.getOpenFileName(self, "选取文件", "C:/", "All Files (*)")
        if os.path.splitext(fileName)[1]in'.fmd':
            fileName1=fileName
        QApplication.processEvents()

    def keyPressEvent(self, e):
        if e.key() == Qt.Key_Escape:
            self.close()
########################两小窗口分离###############################
    def custom_right_menu(self, pos):
        menu = QMenu()
        opt1 = menu.addAction("设置为第一段起始点")
        opt2 = menu.addAction("设置为第一段终点")
        opt3 = menu.addAction("设置为第二段起始点")
        opt4 = menu.addAction("设置为第二段终点")
        action = menu.exec_(self.mainwindow.centralwidget.mapToGlobal(pos))
        if action == opt1:
        ###################起点初始化
            self.plot.addstart(0)
                # do something
            return
        elif action == opt2:
            self.plot.addend(0)
                # do something
            return
        elif action==opt3:
            self.plot.addstart(1)

            return
        elif action==opt4:
            self.plot.addend(1)

            return
        else:
            return
    def custom_right_menu1(self, pos):
        menu = QMenu()
        opt1 = menu.addAction("设置为第一段起始点")
        opt2 = menu.addAction("设置为第一段终点")
        opt3 = menu.addAction("设置为第二段起始点")
        opt4 = menu.addAction("设置为第二段终点")
        action = menu.exec_(self.mainwindow.centralwidget.mapToGlobal(pos))
        if action == opt1:
        ###################起点初始化
            self.plot2.addstart(0)
                # do something
            return
        elif action == opt2:
            self.plot2.addend(0)
                # do something
            return
        elif action==opt3:
            self.plot2.addstart(1)

            return
        elif action==opt4:
            self.plot2.addend(1)

            return
        else:
            return





class Plot3d(QMainWindow):
   def __init__(self, parent=None):
      super().__init__(parent)    #调用父类构造函数
      #self.setWindowTitle("Demo14_1, GUI中的matplotlib绘图")
## rcParams[]参数设置，以正确显示汉字
      #mpl.rcParams['font.sans-serif']=['KaiTi','SimHei']    #汉字字体
      #mpl.rcParams['font.size']=12     #字体大小
      #mpl.rcParams['axes.unicode_minus'] =False    #正常显示负号
      self.__iniFigure()     #创建绘图系统，初始化窗口
      #self.__drawFigure()    #绘图
      #self.test()

##==========自定义函数=================
   def __iniFigure(self):    ##创建绘图系统，初始化窗口
      self.__fig=mpl.figure.Figure(figsize=(10, 8))      #单位英寸
      #self.__fig.suptitle("plot in GUI application")    #总的图标题
      self.figCanvas = FigureCanvas(self.__fig)    #创建FigureCanvas对象
      self.test()
      # 连接事件
#      self.button_plot.clicked.connect(self.plot_)

      # 设置布局
      #layout = QHBoxLayout(self)
      #self.horizontalLayout_3.addWidget(figCanvas)
      #layout.addWidget(figCanvas)
      #self.setLayout(layout)
      #naviToolbar=NavigationToolbar(figCanvas,self)    #创建工具栏
 #     naviToolbar.setToolButtonStyle(Qt.ToolButtonTextUnderIcon)
#      self.addToolBar(naviToolbar)    #添加工具栏到主窗口
#      self.setCentralWidget(figCanvas)




   def test(self):
       # fig, ax = plt.subplots()
       self.ax = self.__fig.gca(projection='3d')
       # self.ax.elev=0
       # self.ax.azim=10
       # self.ax.proj_type='ortho'
       self.ax.set_xlim3d(-10, 10)
       x_major_locator = MultipleLocator(1)
       # 把x轴的刻度间隔设置为1，并存在变量里
       y_major_locator = MultipleLocator(1)
       z_major_locator = MultipleLocator(1)
       # 把y轴的刻度间隔设置为10，并存在变量里
       # ax = plt.gca()
       # ax为两条坐标轴的实例
       self.ax.xaxis.set_major_locator(x_major_locator)
       # 把x轴的主刻度设置为1的倍数
       self.ax.yaxis.set_major_locator(y_major_locator)
       self.ax.zaxis.set_major_locator(z_major_locator)

       self.ax.set(xlabel='X', ylabel='Y', zlabel='Z')
       # self.ax.set_zticklabels()
       theta = np.linspace(-4 * np.pi, 4 * np.pi, 100)
       self.z = np.linspace(-2, 2, 100)
       r = self.z ** 2 + 1
       self.x = r * np.sin(theta)
       self.y = r * np.cos(theta)
       self.x1 = np.linspace(-10, 10, 100)
       self.y1 = np.linspace(-1, 1, 100)
       self.z1 = np.linspace(5, 10, 100)
       self.x2 = np.linspace(-4, 4, 100)
       self.y2 = np.linspace(-2, 2, 100)
       self.z2 = np.linspace(0, 10, 100)
       # self.y2=np.linspace()
       self.line, = self.ax.plot(self.x, self.y, self.z, picker=3, color='b')
       self.line1, = self.ax.plot(self.x1, self.y1, self.z1, picker=3, color='b')
       self.line2, = self.ax.plot(self.x2, self.y2, self.z2, picker=3, color='b')
       self.flag = np.array([0, 0, 0])
       self.__fig.tight_layout()

       self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
       self.__fig.canvas.mpl_connect('scroll_event', self.call_back)


       #self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
       #self.__fig.canvas.mpl_connect('scroll_event', self.call_back)
       #self.__fig.canvas.mpl_connect('button_press_event', self.call_back)
      #self.__fig.canvas.mpl_disconnect(self.__fig.canvas.manager.key_press_handler_id)
      #self.__fig.savefig('3d_fig.png',dpi=200)

   def call_back(self,event):
      self.axtemp = event.inaxes
      x_min, x_max = self.axtemp.get_xlim3d()
      y_min,y_max=self.axtemp.get_ylim3d()
      z_min, z_max = self.axtemp.get_zlim3d()
      fanwei = (x_max - x_min) / 10
      fanwei_y = (y_max - y_min) / 10
      fanwei_z = (z_max - z_min) / 10

      if event.button == 'up':

          self.axtemp.set(xlim3d=(x_min + fanwei, x_max - fanwei), ylim3d=(y_min + fanwei_y, y_max - fanwei_y),
                     zlim3d=(z_min + fanwei_z, z_max - fanwei_z))
          print('up')

          self.line.remove()
          self.line1.remove()
          self.line2.remove()

          x = self.x
          y = self.y
          z = self.z
          x1 = self.x1
          y1 = self.y1
          z1 = self.z1
          x2 = self.x2
          y2 = self.y2
          z2 = self.z2
          a = list()
          b = list()
          c = list()
          for i in range(100):
              if self.x[i] > x_max - fanwei or self.x[i] < x_min + fanwei:
                  a.append(i)
              elif self.y[i] > y_max - fanwei_y or self.y[i] < y_min + fanwei_y:
                  a.append(i)
              elif self.z[i] > z_max - fanwei_z or self.z[i] < z_min + fanwei_z:
                  a.append(i)
          for i in range(100):
              if self.x1[i] > x_max - fanwei or self.x1[i] < x_min + fanwei:
                  b.append(i)
              elif self.y1[i] > y_max - fanwei_y or self.y1[i] < y_min + fanwei_y:
                  b.append(i)
              elif self.z1[i] > z_max - fanwei_z or self.z1[i] < z_min + fanwei_z:
                  b.append(i)
          for i in range(100):
              if self.x2[i] > x_max - fanwei or self.x2[i] < x_min + fanwei:
                  c.append(i)
              elif self.y2[i] > y_max - fanwei_y or self.y2[i] < y_min + fanwei_y:
                  c.append(i)
              elif self.z2[i] > z_max - fanwei_z or self.z2[i] < z_min + fanwei_z:
                  c.append(i)
          # print(a)

          x = np.delete(x, a, axis=0)
          y = np.delete(y, a, axis=0)
          z = np.delete(z, a, axis=0)
          x1 = np.delete(x1, b, axis=0)
          y1 = np.delete(y1, b, axis=0)
          z1 = np.delete(z1, b, axis=0)
          x2 = np.delete(x2, c, axis=0)
          y2 = np.delete(y2, c, axis=0)
          z2 = np.delete(z2, c, axis=0)

          #self.line.remove()
          #self.line1.remove()
          #self.line2.remove()

          if self.flag[0] == 0:
              self.line, = self.ax.plot(x, y, z, picker=3, color='b')
          elif self.flag[0] == 1:
              self.line, = self.ax.plot(x, y, z, picker=3, color='r')
          if self.flag[1] == 0:
              self.line1, = self.ax.plot(x1, y1, z1, picker=3, color='b')
          elif self.flag[1] == 1:
              self.line1, = self.ax.plot(x1, y1, z1, picker=3, color='r')
          if self.flag[2] == 0:
              self.line2, = self.ax.plot(x2, y2, z2, picker=3, color='b')
          elif self.flag[2] == 1:
              self.line2, = self.ax.plot(x2, y2, z2, picker=3, color='r')
          self.__fig.canvas.draw_idle()

      elif event.button == 'down':

          self.axtemp.set(xlim3d=(x_min - fanwei, x_max + fanwei), ylim3d=(y_min - fanwei_y, y_max + fanwei_y),
                     zlim3d=(z_min - fanwei_z, z_max + fanwei_z))
          print('down')

          self.line.remove()
          self.line1.remove()
          self.line2.remove()
          x = self.x
          y = self.y
          z = self.z
          x1 = self.x1
          y1 = self.y1
          z1 = self.z1
          x2 = self.x2
          y2 = self.y2
          z2 = self.z2
          a = list()
          b = list()
          c = list()
          for i in range(100):
              if self.x[i] > x_max + fanwei or self.x[i] < x_min - fanwei:
                  a.append(i)
              elif self.y[i] > y_max + fanwei_y or self.y[i] < y_min - fanwei_y:
                  a.append(i)
              elif self.z[i] > z_max + fanwei_z or self.z[i] < z_min - fanwei_z:
                  a.append(i)
          for i in range(100):
              if self.x1[i] > x_max + fanwei or self.x1[i] < x_min - fanwei:
                  b.append(i)
              elif self.y1[i] > y_max + fanwei_y or self.y1[i] < y_min - fanwei_y:
                  b.append(i)
              elif self.z1[i] > z_max + fanwei_z or self.z1[i] < z_min - fanwei_z:
                  b \
                      .append(i)
          for i in range(100):
              if self.x2[i] > x_max + fanwei or self.x2[i] < x_min - fanwei:
                  c.append(i)
              elif self.y2[i] > y_max + fanwei_y or self.y2[i] < y_min - fanwei_y:
                  c.append(i)
              elif self.z2[i] > z_max + fanwei_z or self.z2[i] < z_min - fanwei_z:
                  c.append(i)
          # print(a)

          x = np.delete(x, a, axis=0)
          y = np.delete(y, a, axis=0)
          z = np.delete(z, a, axis=0)
          x1 = np.delete(x1, b, axis=0)
          y1 = np.delete(y1, b, axis=0)
          z1 = np.delete(z1, b, axis=0)
          x2 = np.delete(x2, c, axis=0)
          y2 = np.delete(y2, c, axis=0)
          z2 = np.delete(z2, c, axis=0)

          if self.flag[0] == 0:
              self.line, = self.ax.plot(x, y, z, picker=3, color='b')
          elif self.flag[0] == 1:
              self.line, = self.ax.plot(x, y, z, picker=3, color='r')
          if self.flag[1] == 0:
              self.line1, = self.ax.plot(x1, y1, z1, picker=3, color='b')
          elif self.flag[1] == 1:
              self.line1, = self.ax.plot(x1, y1, z1, picker=3, color='r')
          if self.flag[2] == 0:
              self.line2, = self.ax.plot(x2, y2, z2, picker=3, color='b')
          elif self.flag[2] == 1:
              self.line2, = self.ax.plot(x2, y2, z2, picker=3, color='r')
          self.__fig.canvas.draw_idle()

   def onpick2(self, event):

      #print(event.artist==self.line)
      if event.mouseevent.button not in ['up','down']:

          if event.artist == self.line:
              if self.flag[0] == 0:
                  # self.line.remove()
                  # self.line,=self.ax.plot(self.x, self.y, self.z, label='parametric curve', picker=5,color='r')
                  self.line.set_color('r')
                  self.__fig.canvas.draw_idle()
                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.flag[0] = 1
              elif self.flag[0] == 1:
                  # self.line.remove()
                  self.line.set_color('b')
                  # self.line,=self.ax.plot(self.x, self.y, self.z, label='parametric curve', picker=5, color='b')
                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.__fig.canvas.draw_idle()
                  self.flag[0] = 0
          elif event.artist == self.line1:
              if self.flag[1] == 0:
                  # self.line1.remove()
                  # self.line1, = self.ax.plot(self.x1, self.y1, self.z1, picker=5, color='r')
                  self.line1.set_color('r')
                  self.__fig.canvas.draw_idle()

                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.flag[1] = 1
              elif self.flag[1] == 1:
                  # self.line1.remove()
                  # self.line1,=self.ax.plot(self.x1, self.y1, self.z1, picker=5, color='b')
                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.line1.set_color('b')
                  self.__fig.canvas.draw_idle()
                  self.flag[1] = 0
          elif event.artist == self.line2:
              if self.flag[2] == 0:
                  # self.line2.remove()
                  # self.line2, = self.ax.plot(self.x2, self.y2, self.z2, picker=5, color='r')
                  self.line2.set_color('r')
                  self.__fig.canvas.draw_idle()
                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.flag[2] = 1
              elif self.flag[2] == 1:
                  # self.line2.remove()
                  self.line2.set_color('b')
                  # self.line2,=self.ax.plot(self.x2, self.y2, self.z2, picker=5, color='b')
                  # self.__fig.canvas.mpl_connect('pick_event', self.onpick2)
                  self.__fig.canvas.draw_idle()
                  self.flag[2] = 0
class PlotWidget(QMainWindow):
    def __init__(self,parent=None):
        super().__init__(parent)
        self.initUi()

    def initUi(self):
        self.getindex()

    def plot(self):

        pg.setConfigOption('background', 'w')
        pg.setConfigOption('foreground', 'k')

        self.win= pg.GraphicsWindow()



        self.win.setWindowTitle('信号图')


        self.label = pg.LabelItem(justify='right')
        self.win.addItem(self.label)
        self.p1 = self.win.addPlot(row=1, col=0,title='Y-X')
        self.p2 = self.win.addPlot(row=2, col=0,title='dY/dX-X')
        self.p1.setMouseEnabled(x=True, y=True)
        self.p2.setMouseEnabled(x=True, y=True)
        self.p1.setMenuEnabled(enableMenu=False, enableViewBoxMenu='same')
        self.p1.setMenuEnabled(enableMenu=False, enableViewBoxMenu='same')
        self.p1.setXLink(self.p2)

        #self.label1 = pg.LabelItem(justify='left')
        #self.label1.setText(
         #   "<span style='font-size: 18pt'>ftydurtdr</span>")
        #self.p1.addItem(self.label1)

        self.p1.setAutoVisible(y=True)
        self.getindex()


        detal_s = (self.s[1:np.size(self.s)] - self.s[0:(np.size(self.s) - 1)]) / (self.a[1:np.size(self.a)] - self.a[0:(np.size(self.a) - 1)])
        self.s1=np.insert(detal_s,0,self.s[0])
        #detal_a = self.a[1:np.size(self.a) - 1]
        #self.a1=np.insert(detal_a,0,self.a[0])
        self.p1.plot(self.a,self.s, pen="g")
        self.p2.plot(self.a,self.s1, pen="b")
        self.p2.setLabel("bottom","深度(m)")
        self.p1.setLabel("left",self.label_left1)
        self.p2.setLabel("left", self.label_left2)

        #self.vLine = pg.InfiniteLine(angle=90, movable=True)

        self.vLine1 = pg.InfiniteLine(angle=90, movable=True)
        self.hLine = pg.InfiniteLine(angle=0, movable=False)
        self.vLine1.setPen('k',style=Qt.DashDotDotLine)
        self.hLine1 = pg.InfiniteLine(angle=0, movable=False)
        self.vLine=pg.InfiniteLine(angle=90, movable=True)
        self.vLine.setPen('k', style=Qt.DashDotDotLine)
        self.hLine.setPen('k', style=Qt.DashDotDotLine)
        self.hLine1.setPen('k', style=Qt.DashDotDotLine)

        self.x=self.a[0]
        self.lable_y=self.s[0]
        self.lable_diffy=self.s1[0]

        self.p1.addItem(self.hLine1)
        self.p1.addItem(self.vLine1)
        self.p2.addItem(self.vLine)
        self.p2.addItem(self.hLine)

        self.plot_data_start = self.p1.plot()
        self.plot_data1_start = self.p2.plot()
        self.plot_data_end = self.p1.plot()
        self.plot_data1_end = self.p2.plot()
        self.data_start_x=list()
        self.data_start_y=list()
        self.data_start_diffy = list()
        self.data_end_x=list()
        self.data_end_y=list()
        self.data_end_diffy = list()
        self.count_start=list()
        self.count_end=list()





        #self.vb = self.p1.vb

        #self.proxy1 = pg.SignalProxy(self.p1.scene().sigMouseMoved , rateLimit=60, slot=self.mouseMoved)
        #self.proxy2 = pg.SignalProxy(self.p2.scene().sigMouseMoved, rateLimit=60, slot=self.mouseMoved)
        #self.proxy = pg.SignalProxy(self.p2.scene().sigMouseMoved, rateLimit=60, slot=self.mouseMoved)
        #self.p1.scene().sigMouseMoved.connect(self.mouseMoved)
        #self.proxy3 = pg.SignalProxy(self.vLine.scene().sigPositionChangeFinished, rateLimit=60, slot=self.statusChange)
        self.vLine.sigPositionChangeFinished.connect(self.labelChange)
        self.vLine1.sigPositionChangeFinished.connect(self.labelChange1)
        #self.win.show()
        self.ex = pyqtgraph.exporters.ImageExporter(self.win.scene())
        self.ex.parameters()['width']=2000
        #ex.export(fileName="test.png")

    def getindex(self):
        if not fileName1:
            return
        file=open(fileName1,'r',encoding='utf-8')
        Deep = list()
        Zdata = list()
        Hdata = list()
        Bdata = list()
        Pitch = list()
        Roll = list()
        Heading = list()

        try:
            self.VerHard= file.readline().strip("\n")  # 硬件版本号，例如19A3表示2019年10(对应A)月份第3版
            self.VerSoft = file.readline().strip("\n")  # 软件版本号，同上
            self.ProjectName = file.readline().strip("\n")  # 工程名称
            self.ConstrucInstitution = file.readline().strip("\n")  # 施工单位
            self.TestInstitution = file.readline().strip("\n")  # 检测执行单位名称
            self.TesterName = file.readline().strip("\n")  # 检测员姓名(编号)
            self.TestDate = file.readline().strip("\n")  # 检测时间
            self.DesignLength = file.readline().strip("\n")  # 钢筋笼设计长度，单位m（四舍五入取整）
            self.TestPointNo = file.readline().strip("\n")  # 测点编号，对应孔桩编号
            self.InitialDepth = file.readline().strip("\n")  # 初始深度，单位cm
            self.PileDiameter = file.readline().strip("\n")  # 基桩的直径，单位cm
            self.Orientation = file.readline().strip("\n")  # 测孔相对桩的方位偏角，单位度
            self.Range = file.readline().strip("\n")  # 桩和测孔的最小距离，单位cm
            self.StepDistance = file.readline().strip("\n")  # 检测步距： 5、10、15、20、15, 单位cm
            self.TestMode = file.readline().strip("\n")  # 采集方式，手动0；自动1
            self.MoveDirection = file.readline().strip("\n")  # 采集方向，从上往下0；从下往上1
            # TestData[]; // 总长度为[(一次采集 * 测试次数] Bytes,// ADC数据按采集时间顺序从前往后依次保存，且高位在前低位在后。
            # 深度值（4个字节） + 磁垂直Z分量+磁水平H分量+磁B总量 +罗盘Pitch俯仰角+ Roll横滚角+ Heading方位角 +0x0D+ 0x0A
            for line in file:
                Deep.append(file.read(4))
                file.read(2)
                Zdata.append(file.readline(6))
                file.read(2)
                Hdata.append(file.readline(6))
                file.read(2)
                Bdata.append(file.readline(6))
                file.read(2)
                Pitch.append(file.read(12))
                file.read(2)
                Roll.append(file.read(12))
                file.read(2)
                Heading.append(file.read(12))


        finally:
            file.close()
            self.a = np.array(Deep, dtype=float)
            self.zdata=np.array(Zdata,dtype=float)
            self.hdata=np.array(Hdata,dtype=float)
            self.bdata=np.array(Bdata,dtype=float)
            self.pitch=np.array(Pitch,dtype=float)
            self.roll=np.array(Roll,dtype=float)
            self.headind=np.array(Heading,dtype=float)

            #print(Zdata)
            #print(Hdata)
            #print(Bdata)
            #print(Pitch)
            #print(Roll)
            #print(Heading)
    def addstart(self,k):
        self.x = self.vLine.getXPos()
        if k==1 and len(self.data_start_x)==0:
            return
        elif self.x<self.a[0]:
            return
        elif len(self.data_start_x)>=k+1:
            self.data_start_x[k]=self.x
            self.data_start_y[k]=self.lable_y
            self.data_start_diffy[k]=self.lable_diffy
            self.count_start[k]=self.count1
            self.plot_data_start.setData(self.data_start_x, self.data_start_y,pen=None, symbol='+')
            self.plot_data1_start.setData(self.data_start_x, self.data_start_diffy,pen=None, symbol='+')
        else:
            self.data_start_x.append(self.x)
            self.data_start_y.append(self.lable_y)
            self.data_start_diffy.append(self.lable_diffy)
            self.count_start.append(self.count1)
            self.plot_data_start.setData(self.data_start_x, self.data_start_y,pen=None, symbol='+')
            self.plot_data1_start.setData(self.data_start_x, self.data_start_diffy,pen=None, symbol='+')
#####################改变曲线该段颜色#############################

    def addend(self,k):
        if k == 1 and len(self.data_end_x) == 0:
            return
        elif self.x<self.a[0]:
            return
        elif len(self.data_end_x) >= k + 1:
            self.data_end_x[k] = self.x
            self.data_end_y[k] = self.lable_y
            self.data_end_diffy[k] = self.lable_diffy
            self.count_end[k]=self.count
            self.plot_data_end.setData(self.data_end_x, self.data_end_y,pen=None, symbol='o')
            self.plot_data1_end.setData(self.data_end_x, self.data_end_diffy,pen=None, symbol='o')
        else:
            self.data_end_x.append(self.x)
            self.data_end_y.append(self.lable_y)
            self.data_end_diffy.append(self.lable_diffy)
            self.count_end.append(self.count)
            self.plot_data_end.setData(self.data_end_x, self.data_end_y,pen=None, symbol='o')
            self.plot_data1_end.setData(self.data_end_x, self.data_end_diffy, pen=None,symbol='o')

    def labelChange(self):
        self.x=self.vLine.getXPos()
        self.vLine1.setPos(self.x)
        b = np.abs(self.a - self.x)
        c = np.where(b == np.min(b))
        index = c[0][0]
        if index>0 and index<(len(self.a)-1):
            if (self.a[index] - self.x) * (self.a[index + 1] - self.x) < 0:
                self.count = index
                self.count1 = index + 1
            else:
                self.count = index - 1
                self.count1 = index
            self.lable_y = (self.s[self.count1] - self.s[self.count]) * (self.x - self.a[self.count]) / (self.a[self.count1] - self.a[self.count]) + \
                      self.s[self.count]
            self.lable_diffy = (self.s1[self.count1] - self.s1[self.count]) * (self.x - self.a[self.count]) / (self.a[self.count1] - self.a[self.count]) + \
                          self.s1[self.count]

            self.label.setText(
                "<span style='font-size: 18pt'>x=%0.2f,   y=%0.2f,   dy/dx=%0.2f</span>" % (
                    self.x, self.lable_y, self.lable_diffy))
    def labelChange1(self):
        self.x = self.vLine1.getXPos()
        self.vLine.setPos(self.x)
        b = np.abs(self.a - self.x)
        c = np.where(b == np.min(b))
        index = c[0][0]
        if index > 0 and index < (len(self.a) - 1):
            if (self.a[index] - self.x) * (self.a[index + 1] - self.x) < 0:
                self.count = index
                self.count1 = index + 1
            else:
                self.count = index - 1
                self.count1 = index
            self.lable_y = (self.s[self.count1] - self.s[self.count]) * (self.x - self.a[self.count]) / (self.a[self.count1] - self.a[self.count]) + \
                      self.s[self.count]
            self.lable_diffy = (self.s1[self.count1] - self.s1[self.count]) * (self.x - self.a[self.count]) / (self.a[self.count1] - self.a[self.count]) + \
                          self.s1[self.count]

            self.label.setText(
                "<span style='font-size: 18pt'>x=%0.2f,   y=%0.2f,   dy/dx=%0.2f</span>" % (
                    self.x, self.lable_y, self.lable_diffy))





if __name__=='__main__':
    QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
    app=QApplication(sys.argv)
    MainWindow = initUI()
    MainWindow.show()
    sys.exit(app.exec_())